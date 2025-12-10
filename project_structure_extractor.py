from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from pathlib import Path

def should_ignore(path, ignore_patterns):
    """Check if path should be ignored based on patterns"""
    path_str = str(path)
    ignore_dirs = {
        'node_modules', '__pycache__', '.git', '.vscode', '.idea',
        'venv', 'env', 'dist', 'build', '.next', 'target'
    }
    ignore_extensions = {'.pyc', '.pyo', '.pyd', '.so', '.dll', '.dylib'}
    
    # Check if any parent directory should be ignored
    for part in Path(path_str).parts:
        if part in ignore_dirs:
            return True
    
    # Check file extension
    if Path(path_str).suffix in ignore_extensions:
        return True
    
    # Check custom patterns
    for pattern in ignore_patterns:
        if pattern in path_str:
            return True
    
    return False

def get_file_language(file_path):
    """Determine programming language based on file extension"""
    ext_map = {
        '.py': 'Python', '.js': 'JavaScript', '.jsx': 'JavaScript React',
        '.ts': 'TypeScript', '.tsx': 'TypeScript React', '.html': 'HTML',
        '.css': 'CSS', '.scss': 'SCSS', '.java': 'Java',
        '.cpp': 'C++', '.c': 'C', '.h': 'C/C++ Header',
        '.go': 'Go', '.rs': 'Rust', '.php': 'PHP',
        '.rb': 'Ruby', '.swift': 'Swift', '.kt': 'Kotlin',
        '.sql': 'SQL', '.sh': 'Shell', '.json': 'JSON',
        '.xml': 'XML', '.yaml': 'YAML', '.yml': 'YAML',
        '.md': 'Markdown', '.txt': 'Text'
    }
    return ext_map.get(Path(file_path).suffix, 'Unknown')

def add_code_block(doc, code, language):
    """Add formatted code block to document"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add light gray background effect by using style
    para.style = doc.styles['Normal']

def export_project_structure(root_dir, output_file, ignore_patterns=None, max_file_size_kb=500):
    """
    Export project structure to Word document
    
    Args:
        root_dir: Root directory of the project
        output_file: Output Word file path
        ignore_patterns: List of patterns to ignore (default: None)
        max_file_size_kb: Maximum file size to include in KB (default: 500)
    """
    if ignore_patterns is None:
        ignore_patterns = []
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Project Structure Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add project info
    doc.add_paragraph(f'Project Root: {os.path.abspath(root_dir)}')
    doc.add_paragraph(f'Generated: {os.popen("date").read().strip()}')
    doc.add_page_break()
    
    # Add table of contents header
    doc.add_heading('Table of Contents', 1)
    toc_para = doc.add_paragraph()
    
    files_data = []
    
    # Walk through directory
    for root, dirs, files in os.walk(root_dir):
        # Filter out ignored directories
        dirs[:] = [d for d in dirs if not should_ignore(os.path.join(root, d), ignore_patterns)]
        
        for file in sorted(files):
            file_path = os.path.join(root, file)
            
            if should_ignore(file_path, ignore_patterns):
                continue
            
            # Check file size
            try:
                size_kb = os.path.getsize(file_path) / 1024
                if size_kb > max_file_size_kb:
                    print(f"Skipping large file: {file_path} ({size_kb:.1f} KB)")
                    continue
            except:
                continue
            
            rel_path = os.path.relpath(file_path, root_dir)
            language = get_file_language(file_path)
            
            files_data.append({
                'path': rel_path,
                'full_path': file_path,
                'language': language
            })
    
    # Add TOC entries
    for idx, file_data in enumerate(files_data, 1):
        toc_para.add_run(f"{idx}. {file_data['path']}\n")
    
    doc.add_page_break()
    
    # Add file contents
    for idx, file_data in enumerate(files_data, 1):
        print(f"Processing: {file_data['path']}")
        
        # Add file header
        heading = doc.add_heading(f"{idx}. {file_data['path']}", 2)
        
        # Add language info
        lang_para = doc.add_paragraph()
        lang_run = lang_para.add_run(f"Language: {file_data['language']}")
        lang_run.italic = True
        lang_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Read and add file content
        try:
            with open(file_data['full_path'], 'r', encoding='utf-8') as f:
                code = f.read()
                add_code_block(doc, code, file_data['language'])
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_data['full_path'], 'r', encoding='latin-1') as f:
                    code = f.read()
                    add_code_block(doc, code, file_data['language'])
            except:
                doc.add_paragraph("[Binary file or encoding error - content not displayed]")
        except Exception as e:
            doc.add_paragraph(f"[Error reading file: {str(e)}]")
        
        doc.add_page_break()
    
    # Save document
    doc.save(output_file)
    print(f"\n✓ Document saved to: {output_file}")
    print(f"✓ Total files processed: {len(files_data)}")

# Example usage
if __name__ == "__main__":
    # Configure your project path here
    PROJECT_ROOT = "."  # Current directory, or specify path like "/path/to/your/project"
    OUTPUT_FILE = "project_structure.docx"
    
    # Add patterns to ignore (optional)
    IGNORE_PATTERNS = [
        '.env',           # Environment files
        '.log',           # Log files
        'package-lock.json',  # Lock files
        '.min.js',        # Minified files
        '.min.css'        # Minified CSS
    ]
    
    # Run the export
    export_project_structure(
        root_dir=PROJECT_ROOT,
        output_file=OUTPUT_FILE,
        ignore_patterns=IGNORE_PATTERNS,
        max_file_size_kb=500  # Skip files larger than 500KB
    )